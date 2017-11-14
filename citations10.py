import scholar
from lxml import html
import lxml.html
from lxml.cssselect import CSSSelector
import requests
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from pprint import pprint
from bs4 import BeautifulSoup

#THIS USES PYHTHON 2.X NOT 3.X

querier = scholar.ScholarQuerier()
settings = scholar.ScholarSettings()
settings._is_configured = True
settings.citform = 4
querier.apply_settings(settings)

user_input = raw_input("Enter the title of one of your citations here: ")
user_input_formatted = '"' + user_input + '"'
print("The article you entered is " + user_input_formatted)

query = scholar.SearchScholarQuery()
#query.set_author("Zach")
# query.set_phrase('"Critique of Methodological Reason"')
query.set_phrase("'" + user_input_formatted + "'")
#query.set_words("computation")
query.set_num_page_results(10)

querier.send_query(query)

#pprint(vars(querier.articles[1]))
#pprint(vars(settings))

# pprint(querier.articles[0].attrs['url_versions'][0])
page = querier.articles[0].attrs['url_versions'][0]
page_request = requests.get(page)
tree = html.fromstring(page_request.content)
authors_time = tree.xpath('//div[@class="gs_a:nth-child(1)"]/text()')
#print(authors_time)

# sel = CSSSelector('div.gs_a:nth-child(2)')
# sel_results = sel(tree)
# print(sel_results)
# match = sel_results[0]
# match_auth_text = match.text
# print(match_auth_text)
# print("the authors of the original article are " + match_auth_text + "\n")

sel2 = CSSSelector('div.gs_fl a:nth-child(4)')
sel_results_2 = sel2(tree)
match2 = sel_results_2[0]
match2_href = match2.get('href')
#print(lxml.html.tostring(match2))
#print(match2_href)


#concatenate the match2_href string with the beginning of the endpoint
#basically need to add this: https://scholar.google.com/ to the front

google_intro = "https://scholar.google.com/"
related_query = google_intro + match2_href
print("the query URL of related articles is " + related_query + "\n")
#related_query is the URL of the "related articles" for the article originally searched

#below is the code that grabs the page of "related articles"
#it returns the text of the gs_a class
related_page_request = requests.get(related_query)
tree2 = html.fromstring(related_page_request.content)
#tree2 is the html tree from the related articles query
#authors_time_2 = tree2.xpath('//div[@class="gs_a"]/text()')


#IF STATEMENT WILL GO HERE (ignore below)
related_author_div = CSSSelector('div.gs_a')
related_author_link_div = CSSSelector('div.gs_a > a')
related_author_div_S = related_author_div(tree2)
related_author_link_div_S = related_author_link_div(tree2)
all_divs = related_author_div_S + related_author_link_div_S
# print(related_author_div_S)
# print(related_author_link_div_S)
# print(all_divs)

title_div = CSSSelector('.gs_rt')
title_div_S = title_div(tree2)

# matchy_match = all_divs[2]
# fucking_match = lxml.html.tostring(matchy_match)

# title_match = title_div_S[2]
# fucking_title_match = lxml.html.tostring(title_match)

# cleantext = BeautifulSoup(fucking_match)
# basetext = cleantext.get_text()
# print(basetext)

# cleantext_title = BeautifulSoup(fucking_title_match)
# basetext_title = cleantext_title.get_text()
# print(basetext_title)

# annoying_shit = CSSSelector('.gs_ctc')
# annoying_shit_2 = CSSSelector('.gs_ctu')
# annoying_shit_array = annoying_shit(tree2)
# annoying_shit_2_array = annoying_shit_2(tree2)
# annoying_shit_combined = annoying_shit_array + annoying_shit_2_array

list_of_annoying_shit = ['[BOOK]', '[CITATION]', '[HTML]', '[BOOK][B]', '[CITATION][C]']
#loop through list of annoying shit to see if any of these are in basetext_title


for x in range(2,10):

	#make strings from all elements

	matchy_match = all_divs[x]
	fucking_match = lxml.html.tostring(matchy_match)

	title_match = title_div_S[x]
	fucking_title_match = lxml.html.tostring(title_match)	

	cleantext_title = BeautifulSoup(fucking_title_match)
	basetext_title = cleantext_title.get_text()
	
	try:
		funk = basetext_title.encode('ascii')
	# print(basetext_title)
	except Exception, e: 
		continue

	cleantext = BeautifulSoup(fucking_match)
	basetext = cleantext.get_text()
	print(basetext)

	if '[BOOK][B]' in funk:
		doob1 = funk.replace('[BOOK][B]', '')
	 	print(doob1)
	elif '[CITATION][C]' in funk:
	  	doob4 = funk.replace('[CITATION][C]', '')
	  	print(doob4)
	elif '[HTML]' in funk:
	  	doob5 = funk.replace('[HTML]', '')
	  	print(doob5)
	elif '[PDF]' in funk:
		doob6 = funk.replace('[PDF]', '')
		print(doob6)
	else:
	  	print(funk)	


	# for y in list_of_annoying_shit:
	# 	if y in basetext_title:
	# 		print("yes")
	# 	 	basetext_title.replace(y, '')

	# print(basetext_title)



#NEED TO CREATE ONE BIG STRING OUT OF ALL THESE TO SEND TO THE WORD DOC


#define the variable to grab the author names from the list of
#related articles
print("----------------------------")

print("Open the word document that has been saved to this folder and copy the text from it, paste it into your paper, and you will have invisibly added 8 citations, which will be indexed in Google Scholar. Overflowing the citation databases will devalue the citation as a commodity and force a reckoning with how scholarship is evaluated today. Thank you!")


# sel_related_authors = CSSSelector('div.gs_a > a')
# sel_related_authors_no_link = CSSSelector('div.gs_a')
# if not sel_related_authors(tree2):
# 	sel_related_authors_result = sel_related_authors_no_link(tree2)
# 	print("nope")
# else:
# 	sel_related_authors_result = sel_related_authors(tree2)
# match_related_authors_first = sel_related_authors_result[0].text
# print("the authors of the first related article are " + match_related_authors_first + "\n")

# for x in range(2, 10):
# 	match_related_authors_first_inc = sel_related_authors_result[x].text
# 	print(match_related_authors_first_inc)



#if it returns a cluster, we have to do something




#print(scholar.citation_export(query2))
#page = requests.get('http://econpy.pythonanywhere.com/ex/001.html')
#tree = html.fromstring(page.content)





#below this line is word doc creator
document = Document()
run = document.add_paragraph().add_run(authors_time)
font = run.font
font.name = 'Calibri'
font.size = Pt(12)
document.save('demo.docx')
