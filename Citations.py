# import Tkinter
# from Tkinter import *
from PyQt5 import QtCore, QtGui, QtWidgets
from mainwindow import Ui_MainWindow
import sys
import time
import traceback
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from pprint import pprint
import requests
import requests
import lxml
import lxml.html
from lxml import html
from lxml.cssselect import CSSSelector
import cssselect
from bs4 import SoupStrainer, BeautifulSoup
import unicodedata
from os.path import expanduser


class StartQT5(QtWidgets.QMainWindow):
    def __init__(self, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        # here connect to slots

        

        #QtCore.QObject.connect(self.ui.run_search,QtCore.SIGNAL("clicked()"), self.run_search_script)
        self.ui.run_search.clicked.connect(self.run_search_script)
        self.ui.search_term.clearFocus()
        self.ui.search_term.returnPressed.connect(self.ui.run_search.click)
    def run_search_script(self):
        user_input = str(self.ui.search_term.text());
        user_input_formatted = '"' + user_input + '"'
        search_results = self.ui.search_results
        search_results.clear()
        search_results.appendPlainText("the article you entered is: " + user_input_formatted)

        wordDoc = ""

        try:

            user_article_URL = 'http://scholar.google.com/scholar?hl=en&q='+user_input_formatted
            r = requests.get(user_article_URL)
            search_results.appendPlainText(user_article_URL)
            tree = lxml.html.fromstring(r.text)

            #this is getting the link to the related articles
            sel2 = CSSSelector('div.gs_fl a:nth-child(4)')
            sel_results_2 = sel2(tree)
            match2 = sel_results_2[0]
            match2_href = match2.get('href')

            #this formats the url of the related articles to the one you cited
            google_intro = "https://scholar.google.com/"
            related_query_page1 = google_intro + match2_href
            #search_results.appendPlainText("the query URL of related articles is " + related_query_page1 + "\n")
            #second page of results below
            related_query_page2 = related_query_page1.replace("q=related:", "start=10&q=related:")
            #search_results.appendPlainText("page 2 of queries is " + related_query_page2 + "\n")

            #below is the code that grabs the content of the first page of "related articles"
            related_page_request = requests.get(related_query_page1)
            tree2 = html.fromstring(related_page_request.content)

            #below is the code that grabs the content of the second page of "related articles"
            related_page2_request = requests.get(related_query_page2)
            tree3 = html.fromstring(related_page2_request.content)

            search_results.appendPlainText("related articles are: \n \n")

            related_author_div = CSSSelector('div.gs_a')
            related_author_link_div = CSSSelector('div.gs_a > a')
            related_author_div_S = related_author_div(tree2)
            related_author_link_div_S = related_author_link_div(tree2)
            all_divs = related_author_div_S + related_author_link_div_S

            title_div = CSSSelector('.gs_rt')
            title_div_S = title_div(tree2)

            list_of_annoying_shit = ['[BOOK]', '[CITATION]', '[HTML]', '[BOOK][B]', '[CITATION][C]']

            document = Document()

            name = QtWidgets.QFileDialog.getSaveFileName(self, 'Save File', "Demo.docx", filter="docx (*.docx *.)")

            for x in range(1,10):

                #make strings from all elements

                match = all_divs[x]
                match_string = lxml.html.tostring(match)
                # search_results.appendPlainText(match_string)

                title_match = title_div_S[x]
                title_match_string = lxml.html.tostring(title_match)
                # search_results.appendPlainText(title_match_string) 

                cleantext_title = BeautifulSoup(title_match_string, "lxml")
                basetext_title = cleantext_title.get_text()
                # search_results.appendPlainText(basetext_title)

                # The try statement below is not working... my temporary workaround is below that
                # try:
                #   encoded_title = remove_control_characters(basetext_title).decode('utf-8').encode('utf-8')
                #   search_results.appendPlainText(encoded_title)
                # except Exception as e: 
                #   continue

                encoded_title = basetext_title

                cleantext = BeautifulSoup(match_string, "lxml")
                basetext = cleantext.get_text()
                search_results.appendPlainText(basetext)

                wordDoc += basetext + "\n"

                if '[BOOK][B]' in encoded_title:
                  book_title = encoded_title.replace('[BOOK][B]', '')
                  search_results.appendPlainText(book_title + "\n")
                  wordDoc += book_title + "\n"
                elif '[CITATION][C]' in encoded_title:
                  citation_title = encoded_title.replace('[CITATION][C]', '')
                  search_results.appendPlainText(citation_title + "\n")
                  wordDoc += citation_title + "\n"
                elif '[HTML]' in encoded_title:
                  html_title = encoded_title.replace('[HTML]', '')
                  search_results.appendPlainText(html_title + "\n")
                  wordDoc += html_title + "\n"
                elif '[PDF]' in encoded_title:
                  pdf_title = encoded_title.replace('[PDF]', '')
                  search_results.appendPlainText(pdf_title + "\n")
                  wordDoc += pdf_title + "\n"
                else:
                  search_results.appendPlainText(encoded_title + "\n")
                  wordDoc += encoded_title + "\n\n"

            saveToWordDoc(document, wordDoc)
            try:
                document.save(str(name[0]))
            except:
                e = sys.exc_info()
                print("Unexpected error:", sys.exc_info()[1])
                search_results.appendPlainText("Could not save document: "+str(e[1]))

        except IndexError:
            search_results.appendPlainText("IndexError")
            if not tree2:
                search_results.appendPlainText("Sorry, no results")
                # HEY GABOOSH HELP -- need to figure out new error handler because scholar was doing it by number of articles
            else:
                e = sys.exc_info()
                search_results.appendPlainText(str(e[0]) + "\n")
                search_results.appendPlainText(str(e[1]) + "\n")
                t = ''.join(traceback.format_tb(e[2]))
                search_results.appendPlainText(t + "\n")

        except:
            e = sys.exc_info()
            search_results.appendPlainText(str(e[0]) + "\n")
            search_results.appendPlainText(str(e[1]) + "\n")
            t = ''.join(traceback.format_tb(e[2]))
            search_results.appendPlainText(t + "\n")




        search_results.appendPlainText("----------------------------" + "\n")

        search_results.appendPlainText("Open the word document that has been saved to this folder and copy the text from it, paste it into your paper, and you will have invisibly added 8 citations, which will be indexed in Google Scholar. Overflowing the citation databases will devalue the citation as a commodity and force a reckoning with how scholarship is evaluated today. Thank you! \n \n")



def saveToWordDoc(document, docText):
  run = document.add_paragraph().add_run(docText)
  font = run.font
  font.name = 'Calibri'
  font.size = Pt(4)
  font.color.rgb = RGBColor(0xff, 0xff, 0xff)


def remove_control_characters(s):
    return "".join(ch for ch in s if unicodedata.category(ch)[0]!="C")

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    myapp = StartQT5()
    myapp.show()
    sys.exit(app.exec_())